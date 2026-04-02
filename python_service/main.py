import os
import io
import json
import base64
import logging
import warnings
import datetime
from typing import Optional, Dict, Any, List

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

import pdfplumber
from openai import OpenAI
from dotenv import load_dotenv
from PIL import Image
import threading
import concurrent.futures
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
from rank_bm25 import BM25Okapi
import jieba

# Global Resources
chroma_client = None
collection = None
embedding_model = None
bm25_index = None
bm25_corpus_ids = []
bm25_corpus_texts = []

def get_embedding_model():
    global embedding_model
    if embedding_model is None:
        print("Loading BGE-large-zh model from local path...")
        try:
            # Load from local directory: ./pretrain_model
            embedding_model = SentenceTransformer('./bge-large-zh-v1.5')
        except Exception as e:
            print(f"Error loading model: {e}")
            # Fallback or re-raise
            raise e
    return embedding_model

def initialize_chroma():
    global chroma_client, collection
    print("Initializing ChromaDB...")
    try:
        chroma_client = chromadb.PersistentClient(path="./chroma_db")
        collection = chroma_client.get_or_create_collection(name="material_docs")
    except Exception as e:
        print(f"Error initializing ChromaDB: {e}")

def initialize_bm25():
    global bm25_index, bm25_corpus_ids, bm25_corpus_texts
    print("Initializing BM25 Index...")
    if collection is None:
        print("Collection is None, skipping BM25 init.")
        return
    
    try:
        # Fetch all documents from Chroma
        results = collection.get()
        
        documents = results['documents']
        ids = results['ids']
        
        if not documents:
            print("No documents found for BM25.")
            bm25_corpus_ids = []
            bm25_corpus_texts = []
            bm25_index = None
            return

        bm25_corpus_ids = ids
        bm25_corpus_texts = documents
        
        print(f"Building BM25 index for {len(documents)} documents...")
        tokenized_corpus = [list(jieba.cut_for_search(doc)) for doc in documents]
        bm25_index = BM25Okapi(tokenized_corpus)
        print("BM25 Index built successfully.")
    except Exception as e:
        print(f"Error initializing BM25: {e}")

# 从父目录加载 .env.local 环境变量
load_dotenv(dotenv_path="../.env.local")

# 屏蔽警告
warnings.filterwarnings("ignore")

app = FastAPI()

@app.on_event("startup")
async def startup_event():
    initialize_chroma()
    initialize_bm25()
    # Pre-load embedding model in a separate thread to avoid blocking startup
    threading.Thread(target=get_embedding_model, daemon=True).start()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

api_key = os.getenv("OPENAI_API_KEY")
base_url = os.getenv("OPENAI_BASE_URL")

if not api_key:
    print("警告: 环境变量中未找到 OPENAI_API_KEY")

client = OpenAI(
    api_key=api_key,
    base_url=base_url,
)

def encode_image_base64(image_bytes: bytes) -> str:
    """将图片字节流转换为 base64 字符串"""
    return base64.b64encode(image_bytes).decode('utf-8')

def analyze_image_with_qwen(base64_image: str, doc_type: str) -> Dict[str, Any]:
    """调用 Qwen-VL-Plus 进行端到端分析"""
    
    # 根据 doc_type 构建 Prompt
    prompt_text = f"""你是一个智能助手，负责从图片中提取结构化数据。
    图片类型是: {doc_type}。
    请识别图片内容，并提取以下字段，仅以 JSON 格式返回，不要包含 markdown 标记。
    
    如果类型是 LICENSE (营业执照):
    - 生产厂家 (String)
    - 营业期限 (String）-三种格式：长期/自YYYY-MM-DD至长期（在只有一个成立日期的情况下）/自YYYY-MM-DD至YYYY-MM-DD

    如果类型是 ISO_QUALITY/ISO_SAFETY/ISO_ENV (ISO体系认证):
    - 生产厂家 (String)
    - 证书有效期 (String, 格式 自YYYY-MM-DD至YYYY-MM-DD)

    如果类型是 CERTIFICATE (产品合格证):
    - 生产厂家 (String, 尽可能提取全称)
    - 产品名称/样品名称 (String)
    - 规格型号 (String)
    
    如果类型是 TYPE_REPORT (型式检验报告):
    - 委托单位 (String)
    - 产品名称/样品名称 (String)
    - 规格型号 (String)
    - 报告日期/签发日期 (String, 格式 YYYY-MM-DD)
    
    如果类型是 COMPANY_ACHIEVEMENT (公司业绩):
    - 生产厂家 (String, 对应厂家名称)
    - 国别 (String, 对应国别)
    
    如果类型是 OTHER (其他):
    - 生产厂家 (String, 对应厂家名称)
    
    注意：请首先判断当前页面是否为包含【委托单位、样品名称、检验结论、报告编号】等关键信息的报告首页或结果页。
    """

    try:
        print(f"正在调用 Qwen-VL-Plus 处理 {doc_type}...")
        # 移除可能的空格
        if api_key:
            client.api_key = api_key.strip()
        if base_url:
            client.base_url = base_url.strip()

        response = client.chat.completions.with_raw_response.create(
            model="qwen_vl_max_public",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}"}},
                        {"type": "text", "text": prompt_text},
                    ]
                }
            ]
        )
        
        # 解析原始 JSON 以处理自定义网关封装
        resp_json = json.loads(response.text)
        print(f"API 原始响应结构 keys: {list(resp_json.keys())}")
        
        content = ""
        # 检查是否为封装格式 {"data": {"choices": [...]}, "choices": null}
        if resp_json.get("choices") is None and resp_json.get("data"):
            print("检测到自定义网关封装格式，正在提取 data 字段...")
            data_field = resp_json["data"]
            if data_field.get("choices") and len(data_field["choices"]) > 0:
                content = data_field["choices"][0]["message"]["content"]
            else:
                print("Error: data 字段中未找到 choices")
                return {}
        else:
            # 标准格式
            parsed = response.parse()
            if parsed.choices and len(parsed.choices) > 0:
                content = parsed.choices[0].message.content
            else:
                print("Error: 标准响应中未找到 choices")
                return {}

        print(f"Qwen-VL 提取内容: {content[:100]}...")
        
        # 清理并解析 JSON
        content = content.replace("```json", "").replace("```", "").strip()
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1:
            return json.loads(content[start:end+1])
        return {}
        
    except Exception as e:
        print(f"Qwen-VL 调用失败: {e}")
        return {}

# --- New Functionality: Full Text Extraction & Indexing ---

def extract_full_text_with_qwen(base64_image: str) -> str:
    """调用 Qwen-VL-Plus 提取全文"""
    prompt_text = """请详细识别图片中的所有文字内容，输出自然语言。
要求：
1. 表格内容结合行列标题输出，保持表格结构逻辑清晰
2. 不同行之间换行输出
3. 保持原始文本的段落和格式
4. 对于印章、签名等特殊区域，明确标注其内容和位置
5. 确保数字、日期、公司名称等关键信息准确无误
6. 如果文字模糊不清，请标注"无法识别"并描述大致内容
"""
    
    try:
        print(f"正在调用 Qwen-VL-Plus 进行全文提取...")
        if api_key:
            client.api_key = api_key.strip()
        if base_url:
            client.base_url = base_url.strip()

        response = client.chat.completions.with_raw_response.create(
            model="qwen_vl_max_public",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}"}},
                        {"type": "text", "text": prompt_text},
                    ]
                }
            ]
        )
        
        # Parse logic similar to analyze_image_with_qwen but expecting text
        resp_json = json.loads(response.text)
        content = ""
        
        if resp_json.get("choices") is None and resp_json.get("data"):
             data_field = resp_json["data"]
             if data_field.get("choices") and len(data_field["choices"]) > 0:
                 content = data_field["choices"][0]["message"]["content"]
        else:
            parsed = response.parse()
            if parsed.choices and len(parsed.choices) > 0:
                content = parsed.choices[0].message.content

        # Clean up markdown code blocks if present
        content = content.replace("```markdown", "").replace("```", "").strip()
        return content

    except Exception as e:
        print(f"Qwen-VL 全文提取失败: {e}")
        return ""

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """简单切片"""
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += (chunk_size - overlap)
    return chunks

@app.post("/index")
async def index_document(
    file: UploadFile = File(...),
    documentId: str = Form(...),
    metadata: str = Form(None)
):
    """
    处理文件 -> Qwen提取全文 -> 切片 -> 向量化 -> 存入ChromaDB -> 更新BM25
    """
    print(f"开始索引文档: {documentId}, 文件名: {file.filename}")
    
    # 解析元数据字符串
    meta_header = ""
    if metadata:
        try:
            meta_obj = json.loads(metadata)
            parts = []
            # 优先提取关键字段
            if meta_obj.get("materialName"):
                parts.append(f"物料名称 : {meta_obj['materialName']}")
            if meta_obj.get("manufacturerName"):
                parts.append(f"厂家名称 : {meta_obj['manufacturerName']}")
            if meta_obj.get("model"):
                parts.append(f"规格型号 : {meta_obj['model']}")
            
            if parts:
                # 使用分号和空格分隔，利于分词
                meta_header = "【文档元数据】 " + " ; ".join(parts) + "\n"
                print(f"生成元数据头: {meta_header.strip()}")
        except Exception as e:
            print(f"元数据解析失败: {e}")

    content = await file.read()
    filename = file.filename.lower()
    
    full_text = ""
    
    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                # 1. 预处理：将所有页面转换为图像对象（串行，快速）
                # 注意：pdfplumber 对象不是线程安全的，所以我们先在主线程提取图片
                print(f"PDF 共 {len(pdf.pages)} 页，正在提取图片...")
                page_images = []
                for i, page in enumerate(pdf.pages):
                    # resolution=300 保证清晰度
                    image = page.to_image(resolution=300).original
                    page_images.append((i, image))
                
                print("图片提取完成，开始并行 OCR 解析...")

                # 2. 定义并发任务函数
                def process_page_ocr(page_index, image_obj):
                    try:
                        img_byte_arr = io.BytesIO()
                        image_obj.save(img_byte_arr, format='PNG')
                        base64_img = encode_image_base64(img_byte_arr.getvalue())
                        # 调用 LLM/OCR
                        text = extract_full_text_with_qwen(base64_img)
                        return page_index, text
                    except Exception as e:
                        print(f"Page {page_index+1} OCR Error: {e}")
                        return page_index, ""

                # 3. 并行执行 OCR
                # 使用 ThreadPoolExecutor，最大并发数建议 5-10，避免触发 API 限流
                ocr_results = {}
                with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
                    future_to_page = {
                        executor.submit(process_page_ocr, idx, img): idx 
                        for idx, img in page_images
                    }
                    
                    for future in concurrent.futures.as_completed(future_to_page):
                        idx = future_to_page[future]
                        try:
                            _, text = future.result()
                            ocr_results[idx] = text
                            print(f"Page {idx+1} OCR completed.")
                        except Exception as e:
                            print(f"Page {idx+1} Task Exception: {e}")
                            ocr_results[idx] = ""

                # 4. 按顺序组装全文
                pages_content = []
                for i in range(len(page_images)):
                    page_text = ocr_results.get(i, "")
                    full_text += f"\n\n# Page {i+1}\n\n{page_text}"
                    pages_content.append(page_text)
                    
        elif filename.endswith(".docx") or filename.endswith(".doc"):
            # Word 处理逻辑：直接提取文本
            print("正在处理 Word 文档索引...")
            doc = docx.Document(io.BytesIO(content))
            
            # 提取文本
            doc_text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    doc_text_parts.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        doc_text_parts.append(" | ".join(row_cells))
            
            full_text = "\n".join(doc_text_parts)
            pages_content = [full_text] # Word 视为单一长文本
            
        else:
            # Image
            base64_img = encode_image_base64(content)
            full_text = extract_full_text_with_qwen(base64_img)
            pages_content = [full_text] # Image treated as single page
            
        if not full_text.strip():
            return {"success": False, "message": "No text extracted"}

        # Chunking Strategy: Page-based with overflow handling
        chunks = []
        for page_text in pages_content:
            if not page_text.strip():
                continue
            
            # If page is small enough, keep as one chunk
            if len(page_text) <= 500:
                chunks.append(meta_header + page_text)
            else:
                # Split large page with overlap
                sub_chunks = chunk_text(page_text, chunk_size=500, overlap=150)
                chunks.extend([meta_header + c for c in sub_chunks])
                
        print(f"Generated {len(chunks)} chunks.")

        # Embedding
        model = get_embedding_model()
        embeddings = model.encode(chunks).tolist()

        # Prepare for Chroma
        ids = [f"{documentId}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [{"documentId": documentId, "chunkIndex": i} for i in range(len(chunks))]

        # Add to Chroma
        collection.add(
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadatas,
            ids=ids
        )
        
        # Update BM25 (Re-initialize for simplicity or incremental update)
        # Incremental update is tricky with rank_bm25. We will just add to our in-memory lists and re-create object if needed
        # But for thread safety and simplicity, we'll re-trigger initialization or just append
        # To be safe: Re-initialize in background
        threading.Thread(target=initialize_bm25, daemon=True).start()

        return {"success": True, "chunks_count": len(chunks)}

    except Exception as e:
        print(f"Indexing error: {e}")
        return {"success": False, "message": str(e)}

@app.post("/delete")
async def delete_document(documentId: str = Form(...)):
    """
    删除文档：从 ChromaDB 和 BM25 索引中移除
    """
    global bm25_index, bm25_corpus_ids, bm25_corpus_texts
    print(f"收到删除请求: {documentId}")
    
    # 1. Delete from ChromaDB
    try:
        # Delete where metadata matches documentId
        # Note: In index_document, metadata key is 'documentId'
        collection.delete(where={"documentId": documentId})
        print(f"ChromaDB 删除成功: {documentId}")
    except Exception as e:
        print(f"ChromaDB 删除失败: {e}")
        # Continue to try to clean up memory
        
    # 2. Delete from In-Memory BM25 Lists
    # Filter out IDs that start with the documentId
    # IDs are format: {documentId}_chunk_{i}
    
    if not bm25_corpus_ids:
        return {"success": True, "message": "BM25 was empty"}

    indices_to_remove = []
    prefix = f"{documentId}_"
    
    for i, cid in enumerate(bm25_corpus_ids):
        if cid == documentId or cid.startswith(prefix):
            indices_to_remove.append(i)
            
    if not indices_to_remove:
        print(f"BM25 中未找到相关文档片段: {documentId}")
        return {"success": True}
        
    print(f"BM25 中将移除 {len(indices_to_remove)} 个片段")
    
    # Remove in reverse order to keep indices valid
    for i in sorted(indices_to_remove, reverse=True):
        del bm25_corpus_ids[i]
        del bm25_corpus_texts[i]
        
    # 3. Rebuild BM25 Index
    # Since we modified the lists, we need to re-instantiate BM25Okapi
    if bm25_corpus_texts:
        try:
            # Re-tokenize is expensive if we do it from scratch, but we have texts.
            # Ideally we should have stored tokenized corpus, but we only have raw texts in bm25_corpus_texts.
            # So we tokenize again.
            print("正在重建 BM25 索引...")
            tokenized_corpus = [list(jieba.cut_for_search(doc)) for doc in bm25_corpus_texts]
            bm25_index = BM25Okapi(tokenized_corpus)
            print("BM25 索引重建完成")
        except Exception as e:
             print(f"BM25 重建失败: {e}")
    else:
        bm25_index = None
        print("BM25 索引已清空")
        
    return {"success": True}

class SearchRequest(BaseModel):
    query: str
    top_k: int = 10         # Final results count (RRF Top K)
    bm25_k: int = 10       # BM25 candidates count
    vector_k: int = 10      # Vector candidates count
    allowed_ids: Optional[List[str]] = None # Filter by Document IDs

def verify_candidate_relevance(candidate: Dict[str, Any], query: str) -> (Dict[str, Any], bool):
    
    check_prompt = f"""
        你是一个检索结果“相关性判定与证据抽取”模块。
        任务：判断【检索片段】中是否存在能够**直接回答用户查询**的信息；
        若存在，仅从原文中抽取对应字段返回。
        ------------------------
        【用户查询】
        {query}
        【检索片段】
        {candidate['text']}
        ------------------------
        【判定规则】

        1. 【可回答性（核心）】
        - 若片段中包含明确的信息、数值或结论，使用户无需依赖其他文档或背景知识即可回答查询，判定为“相关”。
        - 允许读表、字段对齐、单位等价（如：2 小时 = 120 min）。
        - 仅有背景描述、检测行为、报告类型而无答案本身 → 不相关。

        2. 【关键词 / 语义匹配】
        - 片段中需出现查询关键词，或其明确等价表达（同义词、行业写法、缩写）。
        - 同一主题但不能回答该问题 → 不相关。

        3. 【实体一致性】
        - 若查询包含特定实体（公司 / 产品等），片段中必须出现该实体或其明确指代。
        - 出现其他实体的信息 → 不相关。

        4. 【禁止推测】
        - 仅依据片段中实际出现的内容判断与抽取。
        - 不确定是否能直接回答 → 不相关。
        ------------------------
        【输出格式（必须严格遵守）】

        - 相关：
        第一行：是
        第二行起：仅抽取与查询直接相关的原文字段，并用 **加粗** 标出
        - 不相关：
        仅输出：否

        禁止任何额外说明或解释。
        ------------------------
        【示例】

        【示例 1】
        用户查询：
        耐火极限是多少？

        检索片段：
        墙体耐火性能检测结果：耐火极限 ≥ 2.0 h。

        输出：
        是
        **耐火极限 ≥ 2.0 h**

        【示例 2】
        用户查询：
        甲醛释放量检测结果

        检索片段：
        检测项目：甲醛释放量
        检测结果：0.06 mg/m³

        输出：
        是
        **甲醛释放量 0.06 mg/m³**

        【示例 3】
        用户查询：
        耐火极限是多少？

        检索片段：
        本报告对墙体耐火性能进行了检测，方法符合相关标准。

        输出：
        否

        【示例 4】
        用户查询：
        海加公司耐火极限检测结果

        检索片段：
        受检单位：海天新材料有限公司
        耐火极限 ≥ 2.0 h

        输出：
        否
        """

    
    try:
        # 使用用户指定的配置创建专用客户端
        rerank_client = OpenAI(
            api_key="e2030d79bb3bfd45cd21762500036082",
            base_url="https://jcpt-open.cscec.com/aijsxmywyapi/0510250001/v1.0/qwen3_32b_public"
        )

        completion = rerank_client.chat.completions.create(
            model="qwen3_32b_public",
            messages=[
                {"role": "user", "content": check_prompt}
            ],
            stream=True,
            extra_body={"enable_thinking": False},
            temperature=0.1,
            max_tokens=2000
        )
        
        # 处理流式响应
        answer = ""
        for chunk in completion:
            if chunk.choices and chunk.choices[0].delta.content:
                answer += chunk.choices[0].delta.content
        
        answer = answer.strip()
        
        if "是" in answer:
            # 提取 LLM 返回的解释（可能包含高亮字段）
            # 假设 LLM 回复格式为："是，相关字段：**XXX**" 或类似
            candidate['llm_reasoning'] = answer
            print(answer)
            return candidate, True
        else:
            return candidate, False
    except Exception as e:
        print(f"LLM Re-rank error for candidate {candidate.get('id')}: {e}")
        # If LLM fails, keep it to be safe
        return candidate, True

@app.post("/search_advanced")
async def search_advanced_endpoint(req: SearchRequest):
    """
    混合检索: BM25 + Vector -> LLM Rerank (Parallelized)
    """
    query = req.query
    top_k = req.top_k
    bm25_k = req.bm25_k
    vector_k = req.vector_k
    allowed_ids = req.allowed_ids
    
    # Check for empty filter
    if allowed_ids is not None and len(allowed_ids) == 0:
        print("Advanced Search: allowed_ids provided but empty. Returning empty results.")
        return {"success": True, "results": []}

    print(f"Advanced Search Query: {query}, TopK: {top_k}, BM25K: {bm25_k}, VectorK: {vector_k}, FilteredDocs: {len(allowed_ids) if allowed_ids else 'None'}")
    
    bm25_results = []
    vector_results = []

    # Use ThreadPoolExecutor for parallel execution of search and rerank
    with concurrent.futures.ThreadPoolExecutor(max_workers=20) as executor:
        
        # --- Phase 1: Parallel Retrieval (BM25 & Vector) ---
        
        def run_bm25():
            results = []
            if bm25_index:
                # BM25 取 Top K * 2 (增加候选数量以便去重)
                tokenized_query = list(jieba.cut_for_search(query))
                scores = bm25_index.get_scores(tokenized_query)
                
                # Filter indices if allowed_ids is present
                # Since we need sorted scores, we sort indices first
                # Optimization: Only consider indices that are allowed?
                # BM25Okapi doesn't support pre-filtering easily. We post-filter top candidates.
                # To ensure we get enough candidates after filtering, we might need to fetch more.
                # If filter is very restrictive, this might be inefficient.
                # But for now, let's fetch more candidates if filter is present.
                
                fetch_limit = bm25_k * 2
                if allowed_ids:
                    fetch_limit = min(len(bm25_corpus_ids), 1000) # Fetch more to filter
                else:
                    fetch_limit = min(fetch_limit, len(bm25_corpus_ids))
                
                top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:fetch_limit]
                
                # 同文件去重逻辑：只保留每个文件得分最高的片段 -> 移除去重，保留所有高分片段
                # seen_docs = set()
                allowed_set = set(allowed_ids) if allowed_ids else None
                
                for idx in top_indices:
                    # 解析 documentId (格式: {documentId}_chunk_{i})
                    chunk_id = bm25_corpus_ids[idx]
                    # 假设 ID 格式为 "filename_chunk_index"，取最后一个 "_" 之前的部分作为 doc_id
                    if "_chunk_" in chunk_id:
                        doc_id = chunk_id.rsplit("_chunk_", 1)[0]
                    else:
                        doc_id = chunk_id
                    
                    # Filtering
                    if allowed_set and doc_id not in allowed_set:
                        continue
                        
                    # if doc_id in seen_docs:
                    #     continue
                        
                    # seen_docs.add(doc_id)
                    results.append({
                        "id": chunk_id,
                        "text": bm25_corpus_texts[idx],
                        "score": scores[idx],
                        "source": "bm25",
                        "doc_id": doc_id # 记录 doc_id 以便后续使用
                    })
                    
                    # BM25 最终保留 Top K 个片段 (不再限制不同文件)
                    if len(results) >= bm25_k:
                        break
            return results

        def run_vector():
            results = []
            if collection:
                # Vector 取 Top K(增加候选数量以便去重) -> 此时不需要为了去重取更多，但保持逻辑一致
                model = get_embedding_model()
                query_embedding = model.encode([query]).tolist()
                
                where_filter = None
                if allowed_ids:
                    # ChromaDB 'where' clause for filtering
                    # If allowed_ids has 1 item, use equal, else use $in
                    if len(allowed_ids) == 1:
                        where_filter = {"documentId": allowed_ids[0]}
                    else:
                        where_filter = {"documentId": {"$in": allowed_ids}}
                
                search_results = collection.query(
                    query_embeddings=query_embedding,
                    n_results=vector_k,
                    where=where_filter,
                    include=["documents", "metadatas", "distances"]
                )
                
                # seen_docs = set()
                
                if search_results['ids'] and search_results['ids'][0]:
                    for i in range(len(search_results['ids'][0])):
                        chunk_id = search_results['ids'][0][i]
                        metadata = search_results['metadatas'][0][i]
                        
                        # 从 metadata 获取 doc_id，如果没有则解析 ID
                        doc_id = metadata.get('documentId')
                        if not doc_id:
                             if "_chunk_" in chunk_id:
                                doc_id = chunk_id.rsplit("_chunk_", 1)[0]
                             else:
                                doc_id = chunk_id
                                
                        # if doc_id in seen_docs:
                        #     continue
                            
                        # seen_docs.add(doc_id)
                        
                        results.append({
                            "id": chunk_id,
                            "text": search_results['documents'][0][i],
                            "score": search_results['distances'][0][i],
                            "metadata": metadata,
                            "source": "vector",
                            "doc_id": doc_id
                        })
                        
                        # Vector 最终保留 Top K 个片段 (不再限制不同文件)
                        if len(results) >= vector_k:
                            break
            return results

        # Parallel Execution of Retrieval
        future_bm25 = executor.submit(run_bm25)
        future_vector = executor.submit(run_vector)
        
        bm25_results = future_bm25.result()
        vector_results = future_vector.result()

        # --- Phase 2: Weighted Fusion (0.4 BM25 + 0.6 Vector) ---
        
        # Helper for normalization
        def normalize_scores(score_list, reverse=False):
            if not score_list:
                return []
            min_score = min(score_list)
            max_score = max(score_list)
            range_score = max_score - min_score
            
            normalized = []
            for s in score_list:
                if range_score == 0:
                    normalized.append(1.0 if not reverse else 0.0)
                else:
                    norm = (s - min_score) / range_score
                    if reverse:
                        norm = 1.0 - norm
                    normalized.append(norm)
            return normalized

        # 1. Collect all candidates and scores
        # BM25 scores are relevance (higher is better)
        # Vector scores are distances (lower is better, typically L2 in Chroma)
        
        bm25_score_map = {r['id']: r['score'] for r in bm25_results}
        vector_score_map = {r['id']: r['score'] for r in vector_results}
        
        all_candidate_ids = set(bm25_score_map.keys()) | set(vector_score_map.keys())
        
        # Store candidate objects for later retrieval
        candidate_obj_map = {}
        for r in bm25_results:
            candidate_obj_map[r['id']] = r
        for r in vector_results:
            if r['id'] not in candidate_obj_map:
                candidate_obj_map[r['id']] = r
        
        # 2. Normalize BM25 scores
        # We need to normalize across the *entire* potential set, 
        # but here we only have the top K results. Normalizing within the top K is acceptable approximation.
        # If an ID is missing from BM25 results, it effectively has a low score.
        # However, to normalize properly, we should treat missing values as 0 (or min score).
        # But since we only have the retrieved lists, let's normalize the retrieved ones 
        # and assign a default low value (e.g. 0.0) to missing ones after normalization.
        
        bm25_ids = list(bm25_score_map.keys())
        bm25_raw_scores = [bm25_score_map[id] for id in bm25_ids]
        bm25_norm_scores = normalize_scores(bm25_raw_scores, reverse=False)
        bm25_norm_map = dict(zip(bm25_ids, bm25_norm_scores))
        
        vector_ids = list(vector_score_map.keys())
        vector_raw_scores = [vector_score_map[id] for id in vector_ids]
        vector_norm_scores = normalize_scores(vector_raw_scores, reverse=True) # Distance -> Similarity
        vector_norm_map = dict(zip(vector_ids, vector_norm_scores))
        
        # 3. Compute Weighted Score
        # Weight: BM25 = 0.4, Vector = 0.6
        weighted_candidates = []
        
        for cid in all_candidate_ids:
            s_bm25 = bm25_norm_map.get(cid, 0.0)
            s_vector = vector_norm_map.get(cid, 0.0)
            
            # Filter out candidates with zero score in either source
            if s_bm25 <= 1e-9 or s_vector <= 1e-9:
                continue

            final_score = 0.4 * s_bm25 + 0.6 * s_vector
            
            cand_obj = candidate_obj_map[cid]
            # Create a new dict to avoid modifying original results in place if needed
            # but here it's fine to modify or copy
            new_cand = cand_obj.copy()
            new_cand['weighted_score'] = final_score
            new_cand['bm25_norm'] = s_bm25
            new_cand['vector_norm'] = s_vector
            
            weighted_candidates.append(new_cand)
            
        # 4. Sort by Weighted Score Descending
        weighted_candidates.sort(key=lambda x: x['weighted_score'], reverse=True)
        
        candidates_list = weighted_candidates
        print(f"Weighted Fusion Candidates: {len(candidates_list)}")

        # --- Phase 3: Parallel LLM Rerank ---
        
        valid_results = []
        
        # Submit rerank tasks
        future_to_candidate = {
            executor.submit(verify_candidate_relevance, candidate, query): candidate 
            for candidate in candidates_list
        }
        
        # 收集所有 LLM 判定为相关的结果
        llm_debug_results = []
        for future in concurrent.futures.as_completed(future_to_candidate):
            candidate, is_relevant = future.result()
            llm_result_entry = {
                "id": candidate.get("id"),
                "text": candidate.get("text"),
                "is_relevant": is_relevant,
                "llm_reasoning": candidate.get("llm_reasoning", "")
            }
            llm_debug_results.append(llm_result_entry)
            
            if is_relevant:
                candidate['llm_relevant'] = True
                valid_results.append(candidate)

        # 2. 重新排序：LLM 过滤后的结果，按照原始 RRF 分数（隐含在列表顺序中）或重新排序
        # 由于并发执行打乱了顺序，我们需要根据 candidates_list 的原始顺序来恢复顺序
        # candidates_list 是按 RRF 分数降序排列的
        
        final_results = []
        original_order_map = {cand['id']: i for i, cand in enumerate(candidates_list)}
        
        # 按原始 RRF 排名对通过 LLM 的结果进行排序
        valid_results.sort(key=lambda x: original_order_map.get(x['id'], 9999))
        
        # Deduplicate by doc_id (keep highest score)
        deduped_results = []
        seen_doc_ids = set()
        
        for cand in valid_results:
            doc_id = cand.get('doc_id')
            if not doc_id:
                # Fallback extraction if missing
                if "_chunk_" in cand['id']:
                    doc_id = cand['id'].rsplit("_chunk_", 1)[0]
                else:
                    doc_id = cand['id']
            
            if doc_id not in seen_doc_ids:
                seen_doc_ids.add(doc_id)
                deduped_results.append(cand)
        
        # 3. 重新生成连贯排名 (1, 2, 3...)
        # 只取 Top K
        for rank, candidate in enumerate(deduped_results[:top_k]):
            candidate['score'] = rank + 1
            candidate['is_rank'] = True
            final_results.append(candidate)
            
        # --- Save Debug Info to JSON ---
        try:
            debug_data = {
                "query": query,
                "timestamp": datetime.datetime.now().isoformat(),
                "BM25_results": [
                    {
                        "id": r["id"], 
                        "score": r["score"], 
                        "text": r["text"],
                        "doc_id": r.get("doc_id")
                    } for r in bm25_results
                ],
                "Vector_results": [
                    {
                        "id": r["id"], 
                        "score": r["score"], 
                        "text": r["text"],
                        "doc_id": r.get("doc_id")
                    } for r in vector_results
                ],
                "RRF_candidates": [
                    {
                        "id": c["id"], 
                        "text": c["text"],
                        "doc_id": c.get("doc_id"),
                        "bm25_norm": c.get("bm25_norm"),
                        "vector_norm": c.get("vector_norm"),
                        "weighted_score": c.get("weighted_score")
                    } for c in candidates_list
                ],
                "LLM_results": llm_debug_results,
                "Final_output": [
                    {
                        "id": r["id"], 
                        "rank": r.get("score"), 
                        "text": r["text"]
                    } for r in final_results
                ]
            }
            
            # 使用时间戳生成文件名，避免覆盖
            debug_filename = f"search_debug_{int(datetime.datetime.now().timestamp())}.json"
            debug_filepath = os.path.join(os.getcwd(), debug_filename)
            
            with open(debug_filepath, "w", encoding="utf-8") as f:
                json.dump(debug_data, f, ensure_ascii=False, indent=2)
                
            print(f"Search debug info saved to: {debug_filepath}")
            
        except Exception as e:
            print(f"Failed to save debug info: {e}")

        return {"success": True, "results": final_results}

def process_pdf(file_bytes: bytes, doc_type: str) -> Dict[str, Any]:
    """处理 PDF 文件：将第一页转换为图片并调用 Qwen"""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                return {}
            
            # 仅处理第一页 (通常关键信息在第一页)
            print("正在处理 PDF 第一页...")
            page = pdf.pages[0]
            # 转图片, resolution=300 保证清晰度
            image = page.to_image(resolution=300).original
            
            # 转 bytes
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_bytes = img_byte_arr.getvalue()
            
            # 转换为 base64 并调用模型
            base64_img = encode_image_base64(img_bytes)
            return analyze_image_with_qwen(base64_img, doc_type)
            
    except Exception as e:
        print(f"PDF 处理错误: {e}")
        return {}

def process_image(file_bytes: bytes, doc_type: str) -> Dict[str, Any]:
    """处理图片文件"""
    try:
        base64_img = encode_image_base64(file_bytes)
        return analyze_image_with_qwen(base64_img, doc_type)
    except Exception as e:
        print(f"图片处理错误: {e}")
        return {}

import docx

def analyze_text_with_qwen(text_content: str, doc_type: str) -> Dict[str, Any]:
    """调用 Qwen-Plus (文本模型) 进行结构化提取"""
    
    # 截断过长的文本以防超 token (Qwen-Plus context window is large, but safe to limit)
    # 30k chars is usually safe for modern models
    if len(text_content) > 30000:
        text_content = text_content[:30000] + "...(truncated)"

    prompt_text = f"""你是一个智能助手，负责从文档文本中提取结构化数据。
    文档类型是: {doc_type}。
    文档全文内容如下：
    ---------------------
    {text_content}
    ---------------------
    请识别文档内容，并提取以下字段，仅以 JSON 格式返回，不要包含 markdown 标记。
    
    如果类型是 LICENSE (营业执照):
    - 生产厂家 (String)
    - 营业期限 (String）-三种格式：长期/自YYYY-MM-DD至长期（在只有一个成立日期的情况下）/自YYYY-MM-DD至YYYY-MM-DD

    如果类型是 ISO_QUALITY/ISO_SAFETY/ISO_ENV (ISO体系认证):
    - 生产厂家 (String)
    - 证书有效期 (String, 格式 自YYYY-MM-DD至YYYY-MM-DD)

    如果类型是 CERTIFICATE (产品合格证):
    - 生产厂家 (String, 尽可能提取全称)
    - 产品名称/样品名称 (String)
    - 规格型号 (String)
    
    如果类型是 TYPE_REPORT (型式检验报告):
    - 委托单位 (String)
    - 产品名称/样品名称 (String)
    - 规格型号 (String)
    - 报告日期/签发日期 (String, 格式 YYYY-MM-DD)
    
    如果类型是 COMPANY_ACHIEVEMENT (公司业绩):
    - 生产厂家 (String, 对应厂家名称)
    - 国别 (String, 对应国别)
    
    如果类型是 OTHER (其他):
    - 生产厂家 (String, 对应厂家名称)
    
    注意：请首先判断当前页面是否为包含【委托单位、样品名称、检验结论、报告编号】等关键信息的报告首页或结果页。
    """

    try:
        print(f"正在调用 Qwen-Plus 处理文本 (Length: {len(text_content)})...")
        if api_key:
            client.api_key = api_key.strip()
        if base_url:
            client.base_url = base_url.strip()

        # 使用 qwen-plus 或 qwen-max
        response = client.chat.completions.with_raw_response.create(
            model="qwen-plus", 
            messages=[
                {
                    "role": "user",
                    "content": prompt_text
                }
            ]
        )
        
        resp_json = json.loads(response.text)
        content = ""
        
        # 兼容自定义网关格式
        if resp_json.get("choices") is None and resp_json.get("data"):
            data_field = resp_json["data"]
            if data_field.get("choices") and len(data_field["choices"]) > 0:
                content = data_field["choices"][0]["message"]["content"]
        else:
            parsed = response.parse()
            if parsed.choices and len(parsed.choices) > 0:
                content = parsed.choices[0].message.content

        print(f"Qwen-Plus 提取内容: {content[:100]}...")
        
        content = content.replace("```json", "").replace("```", "").strip()
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1:
            return json.loads(content[start:end+1])
        return {}
        
    except Exception as e:
        print(f"Qwen-Plus 调用失败: {e}")
        return {}

def process_word(file_bytes: bytes, doc_type: str) -> Dict[str, Any]:
    """处理 Word 文件：直接提取文本"""
    try:
        # 使用 python-docx 读取内存中的 bytes
        doc = docx.Document(io.BytesIO(file_bytes))
        
        full_text = []
        
        # 1. 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        # 2. 提取表格文本 (简单平铺，用 | 分隔单元格)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    full_text.append(" | ".join(row_cells))
        
        text_content = "\n".join(full_text)
        
        if not text_content.strip():
            print("Word 文档未提取到文本")
            return {}
            
        return analyze_text_with_qwen(text_content, doc_type)

    except Exception as e:
        print(f"Word 处理错误: {e}")
        return {}

@app.post("/parse")
async def parse_document(
    file: UploadFile = File(...),
    docType: str = Form(...)
):
    print(f"接收到文件: {file.filename}, 类型: {docType}")
    
    content = await file.read()
    filename = file.filename.lower()
    
    result = {}
    
    if filename.endswith(".pdf"):
        result = process_pdf(content, docType)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        result = process_word(content, docType)
    else:
        # 假设是图片
        result = process_image(content, docType)
        
    return {"success": True, "data": result, "raw_text": "Processed by Qwen-VL-Plus"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
